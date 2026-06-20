#!/usr/bin/env python3
import os
import sys
import argparse
from pathlib import Path
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Tuple

# Definir TESSDATA_PREFIX para Termux
os.environ['TESSDATA_PREFIX'] = '/data/data/com.termux/files/usr/share/tessdata/'

# Importações opcionais
try:
    import cv2
    HAS_CV2 = True
except ImportError:
    HAS_CV2 = False

try:
    import numpy as np
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False


class PDF2TextExtractor:
    """Conversor de PDF escaneado para DOCX com OCR - Versão Termux"""
    
    def __init__(self, pdf_path: str, output_path: str = None, verbose: bool = False, lang: str = 'por'):
        self.pdf_path = pdf_path
        self.output_path = output_path or pdf_path.replace('.pdf', '_converted.docx')
        self.verbose = verbose
        self.lang = lang
        self.images = []
    
    def log(self, message: str):
        """Print com timestamp"""
        if self.verbose or "✓" in message or "✗" in message or "=" in message:
            print(message)
    
    def convert_pdf_to_images(self) -> List[Image.Image]:
        """Converte PDF em imagens"""
        self.log(f"[1/4] Convertendo PDF em imagens...")
        images = convert_from_path(self.pdf_path, dpi=300)
        self.log(f"✓ {len(images)} páginas convertidas")
        return images
    
    def preprocess_image(self, image: Image.Image) -> Image.Image:
        """Pré-processa imagem para melhorar OCR"""
        if HAS_CV2 and HAS_NUMPY:
            try:
                import cv2
                import numpy as np
                
                img = np.array(image)
                if len(img.shape) == 3:
                    img = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
                
                img = cv2.fastNlMeansDenoising(img, None, 10, 10, 21)
                clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                img = clahe.apply(img)
                img = cv2.adaptiveThreshold(img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                            cv2.THRESH_BINARY, 11, 2)
                return Image.fromarray(img)
            except Exception as e:
                self.log(f"⚠ Processamento OpenCV falhou: {e}, usando PIL")
                return self._preprocess_image_pil(image)
        else:
            return self._preprocess_image_pil(image)
    
    def _preprocess_image_pil(self, image: Image.Image) -> Image.Image:
        """Pré-processamento básico usando PIL"""
        if image.mode != 'L':
            image = image.convert('L')
        
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)
        enhancer = ImageEnhance.Sharpness(image)
        image = enhancer.enhance(2.0)
        
        return image
    
    def perform_ocr(self) -> List[List[dict]]:
        """Realiza OCR em todas as imagens"""
        self.log(f"[2/4] Realizando OCR (idioma: {self.lang})...")
        all_texts = []
        
        for i, image in enumerate(self.images):
            if self.verbose:
                self.log(f"  Página {i+1}/{len(self.images)}...")
            
            processed = self.preprocess_image(image)
            texts = self.extract_text_with_structure(processed)
            all_texts.append(texts)
        
        self.log(f"✓ OCR concluído")
        return all_texts
    
    def extract_text_with_structure(self, image) -> List[dict]:
        """Extrai texto com informações de estrutura"""
        try:
            data = pytesseract.image_to_data(image, lang=self.lang, output_type=pytesseract.Output.DICT)
        except Exception as e:
            self.log(f"⚠ Erro no OCR: {e}")
            return []
        
        paragraphs = []
        current_y = None
        current_paragraph = []
        
        for i, text in enumerate(data['text']):
            if not text.strip():
                continue
            
            y = data['top'][i]
            # Tentar obter confidence, senão usar valor padrão
            conf = int(data.get('confidence', [50] * len(data['text']))[i]) if 'confidence' in data else 50
            h = data['height'][i]
            
            if current_y is not None and abs(y - current_y) > 20:
                if current_paragraph:
                    confs = [item['conf'] for item in current_paragraph]
                    heights = [item['h'] for item in current_paragraph]
                    avg_conf = sum(confs) / len(confs) if confs else 50
                    avg_height = sum(heights) / len(heights) if heights else 10
                    
                    paragraphs.append({
                        'text': ' '.join([item['text'] for item in current_paragraph]),
                        'y_pos': current_y,
                        'confidence': avg_conf,
                        'height': avg_height
                    })
                    current_paragraph = []
            
            current_y = y
            current_paragraph.append({
                'text': text,
                'conf': conf,
                'h': h
            })
        
        if current_paragraph:
            confs = [item['conf'] for item in current_paragraph]
            heights = [item['h'] for item in current_paragraph]
            avg_conf = sum(confs) / len(confs) if confs else 50
            avg_height = sum(heights) / len(heights) if heights else 10
            
            paragraphs.append({
                'text': ' '.join([item['text'] for item in current_paragraph]),
                'y_pos': current_y,
                'confidence': avg_conf,
                'height': avg_height
            })
        
        return paragraphs
    
    def detect_title(self, para_data: dict, avg_height: float) -> bool:
        """Detecta se parágrafo é título"""
        text = para_data['text'].strip()
        is_large = para_data['height'] > avg_height * 1.5
        is_short = len(text.split()) < 10
        return is_large and is_short
    
    def create_document(self, all_texts: List[List[dict]]) -> Document:
        """Cria documento DOCX com formatação"""
        self.log(f"[3/4] Criando documento DOCX...")
        
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        for page_num, paragraphs in enumerate(all_texts):
            if page_num > 0:
                doc.add_page_break()
            
            if not paragraphs:
                continue
                
            heights = [p['height'] for p in paragraphs]
            avg_height = sum(heights) / len(heights) if heights else 10
            
            for para_data in paragraphs:
                text = para_data['text'].strip()
                if not text:
                    continue
                
                para = doc.add_paragraph()
                
                if self.detect_title(para_data, avg_height):
                    para.style = 'Heading 1'
                    para_format = para.paragraph_format
                    para_format.space_before = Pt(12)
                    para_format.space_after = Pt(6)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(6)
                    para_format.line_spacing = 1.15
                
                run = para.add_run(text)
                if para_data['confidence'] < 60:
                    run.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)
        
        self.log(f"✓ Documento criado com sucesso")
        return doc
    
    def save_document(self, doc: Document):
        """Salva documento DOCX"""
        self.log(f"[4/4] Salvando arquivo...")
        doc.save(self.output_path)
        self.log(f"✓ Arquivo salvo: {self.output_path}")
    
    def convert(self) -> bool:
        """Executa conversão completa"""
        self.log(f"\n{'='*50}")
        self.log(f"PDF2TextExtractor - Termux")
        self.log(f"{'='*50}\n")
        self.log(f"Entrada: {self.pdf_path}")
        self.log(f"Saída: {self.output_path}")
        self.log(f"Idioma: {self.lang}\n")
        
        if not os.path.exists(self.pdf_path):
            self.log(f"✗ Arquivo não encontrado: {self.pdf_path}")
            return False
        
        try:
            self.images = self.convert_pdf_to_images()
            all_texts = self.perform_ocr()
            doc = self.create_document(all_texts)
            self.save_document(doc)
            
            self.log(f"\n{'='*50}")
            self.log(f"✓ CONVERSÃO CONCLUÍDA COM SUCESSO!")
            self.log(f"{'='*50}\n")
            return True
        except Exception as e:
            self.log(f"✗ Erro geral: {e}")
            import traceback
            traceback.print_exc()
            return False


def main():
    parser = argparse.ArgumentParser(description='Converte PDF escaneado em DOCX com OCR (Termux)')
    parser.add_argument('pdf_file', help='Arquivo PDF de entrada')
    parser.add_argument('-o', '--output', help='Arquivo DOCX de saída (padrão: mesmo nome + _converted.docx)')
    parser.add_argument('-l', '--lang', default='por', help='Idioma OCR (padrão: por). Use "eng" para inglês')
    parser.add_argument('--verbose', action='store_true', help='Mostrar progresso detalhado')
    
    args = parser.parse_args()
    
    converter = PDF2TextExtractor(args.pdf_file, args.output, args.verbose, args.lang)
    success = converter.convert()
    
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()